"""
full_round_test.py — end-to-end proof on a throwaway document.

1. Build a .docx containing the deliberate ADA error.
2. Upload it to the Company Policies folder in Drive (real file).
3. Run the bot's real correction code (edit_docx_in_place) ADA -> Equality Act 2010.
4. Download it back and assert the text actually changed in Drive.
5. Trash the test doc so it does not clutter the client's folder.
"""
import io
import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

from docx import Document
from googleapiclient.http import MediaIoBaseUpload
from publish_draft import (
    get_services,
    _download_bytes,
    _replace_in_docx_bytes,
    edit_docx_in_place,
    DOCX_MIME,
)

FOLDER_ID = "1UEPmB6aW3n66PulFPFxNFS4I68pLveYh"  # Company Policies folder
WRONG = "Americans with Disabilities Act (ADA)"
RIGHT = "Equality Act 2010"


def build_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("TEST — Accessibility Policy (safe to delete)", level=1)
    doc.add_paragraph(
        "We adhere to relevant accessibility standards such as the "
        "Americans with Disabilities Act (ADA) and Web Content Accessibility "
        "Guidelines (WCAG)."
    )
    doc.add_paragraph("This is a throwaway document created by full_round_test.py.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main():
    drive, _ = get_services()

    # 1 + 2: create the test doc in Drive
    media = MediaIoBaseUpload(io.BytesIO(build_docx_bytes()), mimetype=DOCX_MIME, resumable=False)
    created = drive.files().create(
        body={"name": "AY Policy Bot — ROUND TEST (safe to delete).docx",
              "parents": [FOLDER_ID], "mimeType": DOCX_MIME},
        media_body=media,
        fields="id,webViewLink",
        supportsAllDrives=True,
    ).execute()
    fid = created["id"]
    print(f"Created test doc: {created.get('webViewLink')}")

    # sanity: confirm the wrong text is present before
    before_text = _replace_in_docx_bytes(_download_bytes(drive, fid), WRONG, WRONG)[1]
    print(f"Occurrences of wrong phrase BEFORE: {before_text}")

    # 3: run the REAL correction code path
    count = edit_docx_in_place(drive, fid, WRONG, RIGHT)
    print(f"edit_docx_in_place applied {count} replacement(s)")

    # 4: download again and verify in Drive
    after_bytes = _download_bytes(drive, fid)
    after_doc = Document(io.BytesIO(after_bytes))
    full = "\n".join(p.text for p in after_doc.paragraphs)
    has_right = RIGHT in full
    has_wrong = WRONG in full
    print(f"After: contains '{RIGHT}'? {has_right} | still contains '{WRONG}'? {has_wrong}")

    ok = count > 0 and has_right and not has_wrong
    print("\nRESULT:", "PASS — document corrected in Drive." if ok else "FAIL — see above.")

    # 5: clean up
    drive.files().delete(fileId=fid, supportsAllDrives=True).execute()
    print("Test doc trashed.")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
